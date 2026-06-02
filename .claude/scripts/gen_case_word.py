"""
生成决策科学案例01 Word文档 — 管理层可分发版本
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, indent=False, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer
    return table

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    return p

# ═══════════════════════════════════════
# 封面
# ═══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('决策科学案例库 · 案例01')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('采购决策优化：当直觉不够用的时候')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4a, 0x6f, 0xa5)

doc.add_paragraph()
doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run('中集环科 · 企管部\n2026年5月')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════
# 导语
# ═══════════════════════════════════════
add_heading_styled('关于这个案例库', level=1)

add_para('企业管理中大量的决策问题——采购、排产、库存、定价——本质上是"有限资源、多重约束下的最优选择"。但现实中，这些决策往往依赖个人经验和直觉。当决策变量超过人脑的工作记忆容量时，经验会自动退化为启发式简化（heuristic simplification），导致系统性地偏离全局最优。')

add_para('本案例库的目的不是教管理者学运筹学，而是通过真实业务场景的决策建模演示，帮助管理层建立"约束感知 → 系统优化 → 数据驱动"的决策思维框架。每个案例追求核心逻辑的可理解性和可迁移性，不追求数学严谨性。')

add_para('案例01从公司最熟悉的采购场景切入，以不锈钢板材采购为主线，展示混合整数线性规划（MILP）如何在一个"价格无法谈判"的市场中找到优化空间。')

add_para('建议使用方式：管理层学习会，每次1个案例，15分钟阅读 + 15分钟讨论。', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════
# 一、场景
# ═══════════════════════════════════════
add_heading_styled('一、场景：一个真实的采购困境', level=1)

add_para('你面前是一张不锈钢板采购计划表。标罐车间下个月需要304/2B不锈钢板约85吨，特罐车间需要约30吨，碳罐车间需要约15吨。', indent=True)

add_para('采购经理告诉你："王总，钢厂报价又涨了。现在下单锁定价格，还是等一等？另外，三个车间的量如果合在一起订，可以要非标定尺板，损耗率从20%降到5%，但一次要订够8吨的起订量——碳罐的量单独不够。"', indent=True)

add_para('几个数字在脑子里转：', indent=True)

add_table(
    ['选项', '优势', '风险'],
    [
        ['锁价下单', '规避涨价风险', '占用资金 + 仓储成本'],
        ['合订非标板', '损耗从20%降至5%', '碳罐订单变更则非标板报废'],
        ['分开订标准板', '灵活、随时可调', '20%废料 = 每5吨钢有1吨变废铁'],
        ['等一等再买', '可能等到价格回落', '阀门8-12周lead time不能等'],
    ]
)

add_para('在纸上算了几笔账，每个方案都有道理又都有问题。最后你说："先把标罐的量锁了吧，其他的再看看。"', indent=True)

add_quote('——这个决策对吗？直觉真的够用吗？')

# ═══════════════════════════════════════
# 二、直觉的局限
# ═══════════════════════════════════════
add_heading_styled('二、为什么直觉会撞墙', level=1)

add_para('上面的场景只涉及3个车间、一种物料、一个时点。但真实情况是：', indent=True)

add_table(
    ['维度', '实际规模', '直觉能处理吗？'],
    [
        ['车间/产线', '5个（标罐/特罐/碳罐/气体罐/小件）', '勉强'],
        ['物料种类', '15+种，6大类', '开始吃力'],
        ['计划周期', '52周', '不可能'],
        ['约束条件', 'MOQ、Lead Time、安全库存、预算上限', '彻底超载'],
    ]
)

add_para('更关键的是，这些决策相互缠绕——改变一个变量，连锁影响十几个变量：', indent=True)

add_para('• 本周订一批316L钢板 → 占用预算 → 下周阀门到采购节点时预算可能不够', indent=True)
add_para('• 延迟采购等钢价回落 → 阀门8周lead time不能等 → 生产断线风险', indent=True)
add_para('• 碳罐+标罐聚合订非标板 → 损耗降低了 → 但碳罐ETO订单变更 → 非标板规格作废', indent=True)

add_quote('这就是"约束耦合"：改变一个变量，连锁影响十几个变量。')

add_para('人脑的工作记忆上限是4-7个信息块。而一个5车间×15物料×52周的采购决策，涉及的变量组合以千计。直觉型决策在这个复杂度下不是"不够好"——是系统性地偏向局部最优。这不是能力问题，是生理极限。', indent=True)

# ═══════════════════════════════════════
# 三、核心逻辑
# ═══════════════════════════════════════
add_heading_styled('三、从"拍板"到"建模"：抓住核心逻辑', level=1)

add_para('与其在几十个变量中挣扎，不如退一步问：这个问题的本质是什么？', indent=True)

add_quote('"在满足所有生产需求的前提下，决定什么时间、买什么物料、买多少，使总成本最小。"')

add_para('总成本 = 采购成本 + 持有成本 + 损耗成本 + 紧急采购成本')

add_para('', font_size=6)

add_heading_styled('三个核心问题', level=2)

add_table(
    ['问题', '答案'],
    [
        ['我们控制什么？', '每周期、每种物料"买不买"（0/1）+ "买多少"（数量）'],
        ['我们追求什么？', '年度总成本最低——不是采购单价最低'],
        ['什么限制我们？', '需求必须满足、MOQ必须达标、Lead Time必须等、库存不能负'],
    ]
)

add_heading_styled('公司特性如何嵌入模型', level=2)

add_para('这不是一个泛泛的采购模型。我们的行业特性决定了约束的特殊性：', indent=True)

add_table(
    ['行业特性', '含义', '模型的体现'],
    [
        ['钢材无法议价', '不锈钢板价格跟随伦镍/铬铁市场价，采购方无定价权', '单价是外生参数（市场给定），优化的不是"把价格谈低"，而是"在给定价格下怎么买得更聪明"'],
        ['ETO订单波动', '特种罐订单规格和时点不确定', '需求不是确定值，需要场景分析——"如果需求±20%呢？"'],
        ['阀门长周期', '底阀/安全阀 lead time 8-12周', '模型被迫提前8-12周下单，否则触发紧急采购惩罚'],
        ['非标板MOQ', '钢厂定尺板有最小起订量', '买则≥MOQ，不买则=0——典型的0/1门槛'],
        ['多车间聚合', '五车间同规格物料可合并采购', '跨车间聚合突破MOQ，但需求不确定性随之聚合'],
    ]
)

add_heading_styled('一张图看懂核心权衡', level=2)

add_para('标准板：单价低（¥15/kg）→ 损耗高（~20%废料）→ 实际可用成本 = 15/(1-0.2) = ¥18.75/kg', indent=True)
add_para('非标板：单价高（¥17/kg）→ 损耗低（~5%废料）→ 实际可用成本 = 17/(1-0.05) = ¥17.89/kg', indent=True)
add_para('贵的反而便宜了约5%。但非标板有MOQ门槛——如果凑不够量，库存成本会吃掉省下的损耗。', indent=True)
add_para('模型的作用不是告诉管理者"总选非标板"，而是精确计算出：在什么条件下选非标、什么条件下选标准、跟谁聚合、什么时候买——给出总成本最低的组合方案。', indent=True)

# ═══════════════════════════════════════
# 四、反直觉洞察
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_styled('四、四个反直觉洞察', level=1)

# 洞察1
add_heading_styled('洞察1："钢价不能谈" ≠ "采购没有优化空间"', level=2)

add_para('直觉：钢价是市场价，采购能省什么？不就是比价吗？', bold=True)

add_para('实际情况：在钢材模型的实际求解中，即使价格完全固定（零议价空间），仅通过规格选择 + 采购时机 + 聚合批量三个维度的优化，综合损耗率从15-20%降至6.79%，15种板材年化节省约127万。', indent=True)

add_quote('优化空间不在"单价"，而在"用法"。')

add_para('当品类价格外生时（大宗商品），真正的杠杆在组合决策上——买什么规格、什么时候买、跟谁合在一起买。这三个问题，单靠经验判断没有一个能给出确定性答案。')

# 洞察2
add_heading_styled('洞察2：局部最优加起来 ≠ 全局最优', level=2)

add_para('直觉：每个车间各自算最优采购计划，合起来就是公司最优。', bold=True)

add_para('实际情况：碳罐单独采购非标板——量不够MOQ，只能买标准板，损耗20%。标罐单独采购——量刚好过MOQ，损耗5%。"各算各的"看起来没问题。但合在一起呢？碳罐+标罐聚合采购非标板：两个都是损耗5%。分开算加权平均损耗可能10%+，聚合后直接砍半。', indent=True)

add_quote('分散决策在数学上保证不了全局最优。这不是管理问题，是数学定理。')

add_para('五车间各自独立做采购计划，最大的损失不是某个车间"算错了"，而是聚合效应被组织边界切断了。采购聚合策略的6个跨BU组，正是用来修复这个结构性损失的。')

# 洞察3
add_heading_styled('洞察3：最便宜的物料往往最贵', level=2)

add_para('直觉：选单价最低的物料。', bold=True)

add_table(
    ['选项', '单价', '损耗率', '实际可用成本'],
    [
        ['标准板', '¥15/kg', '20%', '15 ÷ (1-0.2) = ¥18.75/kg'],
        ['非标板', '¥17/kg', '5%', '17 ÷ (1-0.05) = ¥17.89/kg'],
    ]
)

add_para('贵的反而便宜了¥0.86/kg，约5%。但有个陷阱：非标板有MOQ。如果为了凑MOQ多买了用不完的料，库存成本会吃掉所有省下来的损耗。钢材模型最终非标采购占比55.4%，不是100%——剩下44.6%的场景（需求不够MOQ），标准板就是最优。', indent=True)

add_quote('模型告诉你的不是"哪个好"，而是"在什么条件下哪个好"。')

# 洞察4
add_heading_styled('洞察4：库存不是越少越好，也不是越多越好', level=2)

add_para('直觉：库存是万恶之源，能少则少 ← → 不断货最重要，多备点没错。', bold=True)

add_table(
    ['物料类型', 'Lead Time', '模型策略', '逻辑'],
    [
        ['框架型钢', '2-3周', '低库存，高频采购', '补货快，囤着浪费资金'],
        ['罐体钢板', '6-8周', '中等库存，批量采购', '价值高，平衡持有成本与批量经济性'],
        ['阀门组件', '8-12周', '较高安全库存', '断货代价远大于持有成本（紧急采购价格翻倍）'],
    ]
)

add_para('库存策略不是一条原则能概括的。是Lead Time、价值密度、断货代价三个因素交乘出的最优解——而且每种物料的最优策略都不一样。')

# ═══════════════════════════════════════
# 五、管理启示
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_styled('五、给管理层的五个启示', level=1)

add_heading_styled('启示1：从"经验判断"到"约束感知"', level=2)
add_para('经验丰富的管理者能做好决策，但当变量数量超过工作记忆容量，经验会自动退化为启发式简化——用几个经验法则代替全面权衡。好的管理不要求一个人同时处理50个变量，而是建立一个系统来显式地处理这些约束关系。')

add_heading_styled('启示2：从"部门最优"到"系统最优"', level=2)
add_para('碳罐车间和标罐车间各自算出最优采购计划，加起来不如一个跨车间的聚合方案——道理不难懂，但KPI体系常常奖励的是前者。如果一个优化机会天然跨越组织边界，就需要一个跨越组织边界的决策机制来捕捉它。这不是组织架构问题，是决策架构问题。')

add_heading_styled('启示3：方向比精确数字重要', level=2)
add_para('模型给出的"综合损耗率6.79%、年化节省127万"是一个参考基准。更有价值的是优化方向的指示：哪些物料策略明显偏离最优？哪些约束是真正卡脖子的？参数变了结论还稳不稳？方向对了，数字可以逐步逼近。方向错了，算得再精确也是精确的错误。')

add_heading_styled('启示4：模型的价值不在"替代人"，而在"拓展人的认知边界"', level=2)
add_para('MILP模型不会替代采购经理的判断——市场情报、供应商关系、质量风险这些软信息永远需要人。但模型能做一件人做不到的事：在给定假设下，找到数学上严格最优的解，然后告诉你"如果你相信这些数据和假设，这就是最好的做法"。人的价值在于质疑假设、输入软信息、判断模型覆盖不到的风险。模型的价值在于把人从变量的排列组合中解放出来。')

add_heading_styled('启示5：决策科学化的本质——让隐性权衡显性化', level=2)
add_para('每个采购决策背后都有隐含的权衡——价格vs损耗、批量vs库存、确定性vs灵活性。没有模型时，这些权衡发生在决策者的直觉里，不可见、不可复核、不可复制。有了模型，权衡变成可以讨论的东西：你说非标占比应该提高到70%？好，这意味着要新增一个聚合组，我们来看看哪些产品能聚合，代价是什么。', indent=True)

add_quote('从"拍板"到"说理"——这才是决策科学化的真正含义。')

# ═══════════════════════════════════════
# 六、延伸
# ═══════════════════════════════════════
add_heading_styled('六、延伸：这个思路还能用在哪？', level=1)

add_para('同样的"约束优化"框架可以迁移到公司的其他决策场景：')

add_table(
    ['场景', '核心决策', '关键权衡'],
    [
        ['产能分配', '哪个产线做哪个订单', '切换成本 ↔ 交期满足 ↔ 产能利用率'],
        ['物流调度', '从哪发货到哪个客户', '运输成本 ↔ 时效 ↔ 库存可得性'],
        ['人员排班', '多技能工在工位间调配', '人效 ↔ 柔性 ↔ 技能匹配'],
        ['改善项目组合', '多个项目的优先级排列', '投资回报 ↔ 风险 ↔ 资源约束'],
    ]
)

add_para('这些问题的数学结构本质上是一样的：有限资源、多个约束、一个目标、找最优解。不同的只是参数——就像换了数据但求解器不变。')

# ═══════════════════════════════════════
# 附录
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_styled('附录：技术背景', level=1)

add_para('本案例基于实际项目：MILP采购决策优化三模型体系（MVP通用/钢材/罐箱），使用PuLP + CBC开源求解器，在真实业务数据上验证。', indent=True)

add_table(
    ['指标', '数值'],
    [
        ['求解器', 'PuLP + CBC (开源)'],
        ['钢材模型总成本', '¥12,778,965'],
        ['综合损耗率', '6.79%（优化前15-20%）'],
        ['MIP Gap', '< 0.01%'],
        ['年化可量化节省', '约¥127万'],
        ['求解时间', '< 5秒'],
    ]
)

add_para('本案例库由中集环科企管部编制，仅供内部管理层学习讨论使用。', italic=True)
add_para('版本：V1.0 | 日期：2026年5月29日', italic=True)

# ── 保存 ──
out_path = os.path.expanduser(r'~\Documents\Obsidian Vault\26年中集环科工作区\决策科学案例库\案例01-MILP采购决策优化.docx')
doc.save(out_path)
print(f'Word saved: {out_path}')
