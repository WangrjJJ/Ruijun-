#!/usr/bin/env python3
"""Rewrite sections 2 and 3 of PSP报告 with workshop-differentiated, evidence-grounded content."""

from docx import Document
from docx.shared import Pt
from lxml import etree
import os

DOC_PATH = "C:/Users/01455310/Documents/Obsidian Vault/26年中集环科工作区/精益工作区/PSP问题分析报告.docx"
OUT_PATH = "C:/Users/01455310/Documents/Obsidian Vault/26年中集环科工作区/精益工作区/PSP问题分析报告_V6.0.docx"

doc = Document(DOC_PATH)

# ── Locate paragraphs ──
sec2_start = None
sec3_start = None
sec4_start = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '二、第二层' in t or (t.startswith('二、') and '精益' in t):
        sec2_start = i
    elif '三、2026年精益方针' in t or (t.startswith('三、') and '精益方针' in t):
        sec3_start = i
    elif '四、下个时代的竞争力' in t or (t.startswith('四、') and '竞争力' in t):
        sec4_start = i

print(f"sec2_start={sec2_start}, sec3_start={sec3_start}, sec4_start={sec4_start}")

if sec2_start is None or sec3_start is None or sec4_start is None:
    print("ERROR: Could not find section boundaries")
    # Try to find with different patterns
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if len(t) > 3:
            print(f"P{i}: [{p.style.name}] {t[:100]}")
    exit(1)

# ── Helper to make XML elements ──
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_para(text, heading_level=None, bold=False, italic=False):
    """Create a w:p element."""
    p_elem = etree.Element(f'{{{W}}}p')

    if heading_level:
        pPr = etree.SubElement(p_elem, f'{{{W}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle')
        pStyle.set(f'{{{W}}}val', f'Heading{heading_level}')

    r = etree.SubElement(p_elem, f'{{{W}}}r')
    rPr = etree.SubElement(r, f'{{{W}}}rPr')
    if bold:
        b = etree.SubElement(rPr, f'{{{W}}}b')
    if italic:
        i_e = etree.SubElement(rPr, f'{{{W}}}i')

    # Handle newlines by splitting
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            # Add line break
            br = etree.SubElement(r, f'{{{W}}}br')
        t = etree.SubElement(r, f'{{{W}}}t')
        t.text = line
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return p_elem

def make_empty_para():
    return etree.Element(f'{{{W}}}p')

def make_table(headers, rows):
    """Create a simple w:tbl element."""
    tbl = etree.Element(f'{{{W}}}tbl')

    # Table properties
    tblPr = etree.SubElement(tbl, f'{{{W}}}tblPr')
    tblStyle = etree.SubElement(tblPr, f'{{{W}}}tblStyle')
    tblStyle.set(f'{{{W}}}val', 'LightGridAccent1')
    tblW = etree.SubElement(tblPr, f'{{{W}}}tblW')
    tblW.set(f'{{{W}}}w', '5000')
    tblW.set(f'{{{W}}}type', 'pct')

    # Grid
    tblGrid = etree.SubElement(tbl, f'{{{W}}}tblGrid')
    col_w = 9000 // max(len(headers), 1)
    for _ in headers:
        gridCol = etree.SubElement(tblGrid, f'{{{W}}}gridCol')
        gridCol.set(f'{{{W}}}w', str(col_w))

    # Header row
    tr = etree.SubElement(tbl, f'{{{W}}}tr')
    for h in headers:
        tc = etree.SubElement(tr, f'{{{W}}}tc')
        tcPr = etree.SubElement(tc, f'{{{W}}}tcPr')
        tcW = etree.SubElement(tcPr, f'{{{W}}}tcW')
        tcW.set(f'{{{W}}}w', str(col_w))
        tcW.set(f'{{{W}}}type', 'dxa')
        p = etree.SubElement(tc, f'{{{W}}}p')
        r = etree.SubElement(p, f'{{{W}}}r')
        rPr = etree.SubElement(r, f'{{{W}}}rPr')
        b = etree.SubElement(rPr, f'{{{W}}}b')
        sz = etree.SubElement(rPr, f'{{{W}}}sz')
        sz.set(f'{{{W}}}val', '18')  # 9pt
        t = etree.SubElement(r, f'{{{W}}}t')
        t.text = str(h)
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Data rows
    for row in rows:
        tr = etree.SubElement(tbl, f'{{{W}}}tr')
        for val in row:
            tc = etree.SubElement(tr, f'{{{W}}}tc')
            tcPr = etree.SubElement(tc, f'{{{W}}}tcPr')
            tcW = etree.SubElement(tcPr, f'{{{W}}}tcW')
            tcW.set(f'{{{W}}}w', str(col_w))
            tcW.set(f'{{{W}}}type', 'dxa')
            p = etree.SubElement(tc, f'{{{W}}}p')
            r = etree.SubElement(p, f'{{{W}}}r')
            rPr = etree.SubElement(r, f'{{{W}}}rPr')
            sz = etree.SubElement(rPr, f'{{{W}}}sz')
            sz.set(f'{{{W}}}val', '18')  # 9pt
            t = etree.SubElement(r, f'{{{W}}}t')
            t.text = str(val)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return tbl

# ── Remove paragraphs P18 through P29 from XML body ──
body = doc.element.body
paras_to_remove = []
for i in range(sec2_start, sec4_start):
    if i < len(doc.paragraphs):
        p_elem = doc.paragraphs[i]._element
        paras_to_remove.append(p_elem)

# Remove tables 4-8 (0-indexed)
tables_to_remove = []
for ti in [8, 7, 6, 5, 4]:
    if ti < len(doc.tables):
        tables_to_remove.append(doc.tables[ti]._element)

for t_elem in tables_to_remove:
    try:
        body.remove(t_elem)
    except:
        pass

for p_elem in paras_to_remove:
    try:
        body.remove(p_elem)
    except:
        pass

print(f"Removed {len(paras_to_remove)} paragraphs and {len(tables_to_remove)} tables")

# ── Build new content ──
new_elements = []

# Section 2 heading
new_elements.append(make_para('二、第二层：基于现状诊断的精益战场', heading_level=1))
new_elements.append(make_empty_para())

# 2.1
new_elements.append(make_para('2.1 现状诊断：五个车间的精益起点完全不同', heading_level=2))
new_elements.append(make_para('四个战场是战略方向，但五个车间的起点差异决定了"先从哪打、用什么武器"完全不同。标罐在价格战前线流血，特罐有利润窗口但柔性不够，碳罐有多能工文化但未充分利用，气体罐和小件有速赢空间但非当前主战场。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['车间', '成本紧迫性', '核心痛点', '柔性压力', '知识基础', '改善文化就绪度', 'E项目参与'],
    [
        ['标罐', '最高', 'Q1亏损1,581万，价格战最前线，材料利用率是生死线', '低（同质化产品）', '低（自动焊经验在少数人）', '低（摊派式提案）', '深度参与'],
        ['特罐', '中高', '单台80-120h定制工时，设计-制造协同断裂', '最高（特种罐利润窗口打开）', '中（防波板/CVC有积累）', '中（有意愿缺方法）', '部分参与'],
        ['碳罐', '中', '粉末降本是重点，多能工28%→40%已有路径', '中（小批量多品种）', '中（多能工文化）', '最高（天然改善土壤）', '部分参与'],
        ['气体罐', '低', '产品结构简单，仅焊接气体降本', '低', '低', '低', '未参与'],
        ['小件', '中', '制造分散，10种以上部件整合空间大', '低', '低', '低', '未参与'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('关键判断：标罐是成本战的正面战场（必须打），特罐是利润池的钥匙（必须抢），碳罐是改善文化的最佳试验田（必须用），气体罐和小件是速赢储备（有空再做）。', bold=True))
new_elements.append(make_empty_para())

# 2.2
new_elements.append(make_para('2.2 E项目的启示：什么被验证了，什么还没被验证', heading_level=2))
new_elements.append(make_para('E项目是2026年唯一经过管理层公开验证的改善实践（"成效显著"，W18总结会，2026-04-27）。年化收益1,200万，单台收益约¥4,000。在铺开精益之前，必须先理解它证明了什么、没证明什么——否则就会把已验证的成功路径丢掉，去追逐未经验证的理论方案。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['已被E项目验证（可以横展）', '尚未被验证（需要Phase 0试探）'],
    [
        ['BOM优化+采购协同可产出千万级年化收益', '一线员工能自发发现问题并改善（E项目是管理层推动的）'],
        ['跨部门协作（设计+采购+工艺+生产）是降本前提', '标准作业/视频SOP在罐箱制造场景中被一线接受'],
        ['管理层关注+周例会节奏能持续驱动改善', 'Daily Management在班组层面能自主运行'],
        ['材料利用率是成本第一杠杆（单台¥1,020）', 'SMED/换型改善能真正缩短交付周期'],
        ['三工时统一是成本核算的基础', '激励机制能激发改善意愿（目前仍是摊派≥15件/季）'],
        ['管理层"成效显著"评价=精益有政治资本', '精益方法脱离E项目语境后还能否获管理层信任'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('精益策略的第一原则：从E项目已验证的阵地出发，向未验证的方向逐步试探。不另起炉灶，而是给E项目装上"持续改善"的引擎。', bold=True))
new_elements.append(make_empty_para())

# 2.3
new_elements.append(make_para('2.3 战场一：成本竞争力——分车间、分成本项的精准打击', heading_level=2))
new_elements.append(make_para('不是"降低成本"这个模糊目标，而是先回答三个问题：哪个车间？哪项成本？什么杠杆？每个车间的成本结构和改善杠杆完全不同。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['车间', '最大成本杠杆', '依据', '2026年目标', '方法', '当前状态'],
    [
        ['标罐', '材料利用率', 'E项目验证：材料~50%成本，损耗率6.79%每降1pp=显著增利', '损耗率6.79%→<5.5%', 'E项目BOM横展+筒体板取消裁边', '5/6项已完成，1项进行中'],
        ['标罐', '转化工时', '自动焊/机器人已投入但效率未完全释放', '工时-5~10%', '标准作业消除动作浪费', '自动焊试点推广中'],
        ['特罐', '设计-制造协同', '单台80-120h定制工时，设计变更传导失真', 'ECN落地周期-50%', 'ECN标准化（TPI-03）', 'EBOM→PBOM断裂，未启动'],
        ['全车间', '质量成本', '从未量化，返工/报废/索赔是隐藏利润杀手', '首次量化TOP3质量损失', 'Pareto分析（Phase 0第一步）', '数据在质量部Excel中'],
        ['全车间', '辅材/能耗', '低值易耗-5~10%目标已有，但无工位级消耗基线', '建立工位级消耗基线', '原单位管理（MFG-02）', '未启动'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('2026年不做什么：不在全公司推全面质量成本核算（没人力），不做TPM全面铺开（先OEE基线），不做APS系统（先手工看板3个月），不追求所有车间同时推进。', italic=True))
new_elements.append(make_empty_para())

# 2.4
new_elements.append(make_para('2.4 战场二：柔性响应力——先搞清楚"不柔性的代价"', heading_level=2))
new_elements.append(make_para('柔性不是口号。不柔性的具体代价：不敢接特种罐订单→毛利锁死5%（vs特种罐15-20%）；换型时间长→小批量排产困难→交期拉长→客户流失；多能工不足→产线间无法调峰→忙闲不均。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['柔性瓶颈', '当前状态', '为什么先打这里', '2026年目标', '方法'],
    [
        ['特罐换型', '无基线数据，换型时间未知', '特种罐利润窗口正打开，柔性不够=不敢接=利润流失', '建立换型基线+完成1个SMED试点', '视频记录→分析→改善→验证'],
        ['碳罐多能工', '28%→40%（已有季度分解）', '已有文化基础和多能工目标，最可能成功', '多能工≥32%（Q2目标）', 'TWI-JI试点+技能矩阵'],
        ['碳罐单元化', '外购转自制评估中', '碳罐是小批量多品种天然试验田', '完成单元化试点设计', '3P+流动改善'],
        ['OTD可视化', '无端到端交期数据', '不知道在哪断的，就不知道改善什么', '首次度量OTD+识别TOP3断点', '跨职能VSM（TPI-01）'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('2026年不做什么：不在标罐做SMED（同质化产品换型压力小），不在全公司推多能工（先从碳罐做起），不做APS（先手工看板3个月），不在没基线时设换型目标。', italic=True))
new_elements.append(make_empty_para())

# 2.5
new_elements.append(make_para('2.5 战场三：知识基础设施——从一个老师傅、一个瓶颈工序开始', heading_level=2))
new_elements.append(make_para('知识显性化是最容易被高估速度、低估阻力的战场。三个原因：老师傅的诀窍是职业安全感来源——要求"交出来"会遇到隐性抵抗；SOP写出来不用=废纸——日本精益的教训是标准作业必须由操作者自己写才有ownership；视频化是好方向，但"谁来拍、拍什么、怎么用"不解决就是形式主义。'))
new_elements.append(make_empty_para())

new_elements.append(make_para('Phase 0验证路径（不铺开）：', bold=True))
new_elements.append(make_para('1. 找一个人：不是技术最好的，是最愿意教的那个老师傅（碳罐多能工文化更可能找到）'))
new_elements.append(make_para('2. 选一个工序：不是最复杂的，是新人最容易出错的那个瓶颈工序'))
new_elements.append(make_para('3. 拍一段视频：手机10分钟→LLM生成步骤化文字初稿→师傅30分钟review修正→打印贴现场'))
new_elements.append(make_para('4. 问三个问题（2周后）：有人用吗？出错降了吗？师傅觉得有价值吗？'))
new_elements.append(make_para('5. 验证成功→扩展3个工序→再谈TWI-JI体系；失败→损失2.5小时+手机内存'))
new_elements.append(make_empty_para())

new_elements.append(make_para('2026年不做什么：不做全面TWI-JI培训体系（先证明视频SOP有用），不做技能矩阵（等碳罐多能工自然产出数据），不买知识管理系统。', italic=True))
new_elements.append(make_empty_para())

# 2.6
new_elements.append(make_para('2.6 战场四：管理基础设施——DBSO可以等，但三个"微习惯"不能等', heading_level=2))
new_elements.append(make_para('管理基础设施的最大陷阱是"先建机构→再建流程→最后出成果"。正确顺序是相反的——先出成果，再让成果产生对机构的需求。DBSO不应是企管部推动的行政产物，而应是总经理看到改善成果后主动问"谁在做这个？给他一个正式编制"的自然结果。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['机制', '传统做法', '2026年务实做法', '启动条件', '增量时间'],
    [
        ['DBSO', '先挂牌→找人→做事', '王瑞俊先做事→成果可见→总经理追问→顺势挂牌', 'Phase 0产出可见结果后', '每周半天（已在计划）'],
        ['MOR', '新建月度运营回顾会', '现有经营分析会+3分钟成本TOP3看板→跑通3次→升级议程', 'Pareto分析完成后', '3分钟/月'],
        ['Daily Mgmt', '建SDQCG三层会议体系', '晨会加1个问题"昨天有什么浪费/障碍？"→习惯了再扩展', '零门槛，随时可试', '2分钟/天'],
        ['激励', '先发制度→等案例', '先有3个一线Mini-Kaizen→再发布激励制度', '有案例支撑后', '0（Phase 2启动）'],
        ['人才培养', '建黄/绿/黑带体系', '先让视频SOP被使用→再谈TWI→再谈带级', '视频SOP验证成功后', '0（Phase 2启动）'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('核心原则：管理机制的建立顺序=成果可见性的顺序。没成果前不建任何新组织、不开任何新会议、不发任何新制度。', bold=True))
new_elements.append(make_empty_para())

# ── Section 3 ──
new_elements.append(make_para('三、2026年精益方针：有选择的聚焦', heading_level=1))
new_elements.append(make_empty_para())

# 3.1
new_elements.append(make_para('3.1 总体判断：2026不是全面推进年，是破冰验证年', heading_level=2))
new_elements.append(make_para('2026年不能走"全面导入DBS体系"的路线。三个约束决定上限：'))
new_elements.append(make_para('1. 资源约束：唯一能投入较整块时间的人只有王瑞俊。各部门长/车间主任被日常经营占满（每月至多4-8小时），班组长被生产任务绑死（每天至多15分钟）。'))
new_elements.append(make_para('2. 信任约束：管理层和一线都没见过精益在罐箱制造场景的本地化成功案例。没可见成果前铺开任何体系=形式主义=透支精益信誉。'))
new_elements.append(make_para('3. 时间约束：6月启动Phase 0，有效窗口仅6个月（7-12月）。6个月够验证2-3个核心方法的本地可行性，不够铺开24个课题。'))
new_elements.append(make_empty_para())

new_elements.append(make_para('2026年真实目标：用6个月回答一个问题——精益方法在罐箱制造场景中到底能不能产生可见的、一线认可的、管理层愿意继续投资的成果？', bold=True))
new_elements.append(make_empty_para())

# 3.2
new_elements.append(make_para('3.2 四个方针——每个都有明确的"做什么"和"不做什么"', heading_level=2))
new_elements.append(make_empty_para())

# 方针一
new_elements.append(make_para('方针一：接住E项目势头，把"一个项目"变成"一个可复用的模式"', bold=True))
new_elements.append(make_para('E项目是已点燃的火种。精益的第一件事不是另起炉灶，是给这团火加柴——把E项目的BOM优化方法论、跨部门协作节奏、周例会复盘机制抽象为可横展到其他订单的模板。'))

new_elements.append(make_table(
    ['做什么', '不做什么'],
    [
        ['E项目成果横展至≥2个其他订单', '不另开新项目与E项目争夺管理注意力'],
        ['建立月度E项目&精益联动复盘（复用已有周例会）', '不把E项目重新包装成"精益项目"'],
        ['把BOM优化方法论标准化为可横展模板', '不在验证横展完成前追求方法论完美'],
        ['用E项目"成效显著"作为PSP讨论会事实锚点', '不以"精益"为名另搞一套话语体系'],
    ]
))
new_elements.append(make_empty_para())

# 方针二
new_elements.append(make_para('方针二：标罐打成本歼灭战，特罐建柔性桥头堡——分车间、分打法', bold=True))
new_elements.append(make_para('五个车间不能一套打法。标罐的问题是"怎么在亏损中活下来"——材料利用率是唯一核心指标。特罐的问题是"怎么接到高利润订单"——柔性是接单能力的硬约束。碳罐的问题是"怎么把多能工文化放大为改善文化"——这里有最好的土壤。'))

new_elements.append(make_table(
    ['车间', '2026年方针', '唯一关键指标', '为什么是这个指标', '方法'],
    [
        ['标罐', '成本歼灭', '材料利用率（损耗率6.79%→<5.5%）', 'E项目已证明可行，每1pp降耗=显著增利', 'E项目横展+标准作业消除浪费'],
        ['特罐', '柔性桥头堡', '换型时间首次可度量+1个SMED试点', '柔性不够=不敢接特种罐=毛利锁死5%', 'VSM+SMED+单元化试点'],
        ['碳罐', '改善试验田', '第一个一线自主Mini-Kaizen来自碳罐', '多能工文化=最好的改善土壤', '找牢骚班组长+TWI-JI试点'],
        ['气体罐', '速赢储备', '焊接气体-10%（已有目标）', '产品简单，改起来快，积累成功案例', '原单位管理+可视化'],
        ['小件', '速赢储备', '10种以上部件整合（已有目标）', '分散制造整合空间大', '整合+工装改善'],
    ]
))
new_elements.append(make_empty_para())

# 方针三
new_elements.append(make_para('方针三：知识显性化从"一个"开始，不是从"体系"开始', bold=True))
new_elements.append(make_para('知识管理的最大风险不是做得慢，是做了一堆没人用的SOP文档。2026年只验证一件事：一个工序的视频SOP能不能被新人使用、能不能降低出错率。验证成功→TWI-JI/多能工/技能矩阵有了地基。失败→损失一个下午和一部手机的内存。'))

new_elements.append(make_table(
    ['做什么', '不做什么'],
    [
        ['1个老师傅 × 1个瓶颈工序 × 1个视频SOP', '不做全公司SOP覆盖率目标'],
        ['让操作者自己review修正（不是工程师代写）', '不做TWI-JI培训体系（先验证再设计）'],
        ['2周后观察：有人用吗？出错降了吗？师傅认可吗？', '不买知识管理软件'],
        ['验证成功→扩展3个工序；失败→换方法再试', '不做技能矩阵（等碳罐多能工自然产出数据）'],
    ]
))
new_elements.append(make_empty_para())

# 方针四
new_elements.append(make_para('方针四：把改善藏进日常工作，不让人感觉到"多了一件事"', bold=True))
new_elements.append(make_para('人力资源偏紧是最大约束。如果精益被感知为"新增工作量"，就会死于无声抵抗。策略：精益不是新增工作，是让现有工作变得更好的方法——把改善藏进现成会议、现成流程、现成关系。'))

new_elements.append(make_table(
    ['嵌入点', '方式', '增量时间/人', '启动条件'],
    [
        ['月度经营分析会', '加"成本损失TOP3"看板（3分钟），替代部分现有议程', '0分钟（替代而非新增）', 'Pareto分析完成后'],
        ['班组晨会', '加1个问题"昨天有什么浪费/障碍？"', '2分钟/天', '找到认同的班组长即可'],
        ['工艺部周会', '加"这周哪个设计变更在车间走样了？"', '5分钟/周', 'PSP讨论会后'],
        ['王瑞俊个人周计划', '每周半天固定精益时间（Phase 0核心投入）', '4小时/周', '即刻'],
    ]
))
new_elements.append(make_empty_para())

# 3.3
new_elements.append(make_para('3.3 2026年12月成功的样子', heading_level=2))
new_elements.append(make_para('不是"24个课题启动"，不是"DBSO编制到位"，不是"MOR正式运行"。而是下面五件事。每件事都有对应的验证逻辑——做不到说明什么。'))
new_elements.append(make_empty_para())

new_elements.append(make_table(
    ['#', '2026年12月的成功标志', '如果没做到说明什么'],
    [
        ['1', 'E项目模式横展至≥2个其他订单，产生可核算增量收益', 'E项目成功可能是偶然的，精益方法需重新设计'],
        ['2', '≥3个一线Mini-Kaizen成功案例，每个由班组长本人讲述', '改善文化存在系统性障碍，需先解决激励和信任'],
        ['3', '≥1个SMED试点完成，换型时间可度量下降', '罐箱换型柔性可能不是精益能解决的，需重估优先级'],
        ['4', '月度经营会成本TOP3看板运行≥3次，总经理追问≥1次', '管理层无法通过数据议程建立关注，需换触达路径'],
        ['5', '1个视频SOP被一线实际使用，有证据降低新人出错率', '知识显性化阻力>预期，需改变策略（如先推激励再推知识）'],
    ]
))
new_elements.append(make_empty_para())

new_elements.append(make_para('五件事做到→2027年有资格谈DBSO编制/激励/24课题/全面推进。做不到→精益方法在罐箱场景需根本性重新设计——这本身也是有价值的结论。', bold=True))
new_elements.append(make_empty_para())

# ── Insert new elements into body ──
# Find reference element: the last paragraph of section 1 (P17, which is about ETO/现场有大量技巧)
ref_elem = None
for p in doc.paragraphs:
    t = p.text.strip()
    if '半封闭状态' in t or '工艺只能维持通用标准' in t:
        ref_elem = p._element
        break

# Fallback: find empty paragraph before what was section 4
if ref_elem is None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if '四、下个时代的竞争力' in t:
            prev = p._element.getprevious()
            while prev is not None:
                if prev.tag == f'{{{W}}}p':
                    ref_elem = prev
                    break
                prev = prev.getprevious()
            break

if ref_elem is None:
    # Last resort: use first paragraph
    ref_elem = doc.paragraphs[0]._element

print(f"Inserting {len(new_elements)} elements after reference element")

# Insert in reverse order (each goes right after ref_elem)
for elem in reversed(new_elements):
    ref_elem.addnext(elem)

# ── Save ──
doc.save(OUT_PATH)
print(f"Saved to: {OUT_PATH}")
print("Done!")
